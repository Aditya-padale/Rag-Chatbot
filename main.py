import os
import time
import asyncio
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, Request, UploadFile, File, HTTPException
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferWindowMemory
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import pdfplumber
import docx
from PIL import Image
import pytesseract
import shutil
from langchain.schema import Document
from google.api_core.exceptions import ResourceExhausted, GoogleAPICallError
import logging
from typing import Dict, Any
import json


app = FastAPI()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Enable CORS for deployment
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000", 
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        "https://*.vercel.app",  # Your Vercel frontend
        "https://rag-chatbot-frontend.vercel.app",  # Replace with your actual Vercel URL
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Rate limiting variables
last_request_time = 0
min_request_interval = 1  # Minimum 1 second between requests

# 🔐 Load environment variables
load_dotenv()

# 🔐 Set your Gemini API key from .env file
api_key = os.getenv("GOOGLE_API_KEY")
if api_key:
    os.environ["GOOGLE_API_KEY"] = api_key

# 📚 Load and split knowledge base from TXT with FAISS caching
kb_txt_path = "aditya_full_personal_profile.txt"
faiss_path = "faiss_index"
if not os.path.exists(kb_txt_path):
    raise FileNotFoundError("❌ 'aditya_full_personal_profile.txt' not found. Please add your TXT knowledge base in the project directory.")

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()

gemini_embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

if os.path.exists(faiss_path):
    vectorstore = FAISS.load_local(faiss_path, gemini_embeddings, allow_dangerous_deserialization=True)
else:
    kb_text = extract_text_from_txt(kb_txt_path)
    raw_docs = [Document(page_content=kb_text)]
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    docs = splitter.split_documents(raw_docs)
    vectorstore = FAISS.from_documents(docs, gemini_embeddings)
    vectorstore.save_local(faiss_path)

# 🧠 Memory for conversation (limit to last 10 exchanges to prevent memory issues)
memory = ConversationBufferWindowMemory(
    memory_key="chat_history", 
    return_messages=True, 
    k=10  # Keep only last 10 conversation turns
)

# Initialize LLM with timeout and retry settings
def create_gemini_llm():
    return ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        timeout=30,  # 30 second timeout
        max_retries=3,
        temperature=0.7
    )

# 🧠 RAG Chain with memory
qa_chain = ConversationalRetrievalChain.from_llm(
    llm=create_gemini_llm(),
    retriever=vectorstore.as_retriever(search_kwargs={"k": 4}),
    memory=memory,
    verbose=True
)

# Rate limiting function
async def rate_limit():
    global last_request_time
    current_time = time.time()
    time_since_last_request = current_time - last_request_time
    if time_since_last_request < min_request_interval:
        await asyncio.sleep(min_request_interval - time_since_last_request)
    last_request_time = time.time()

class UserMessage(BaseModel):
    user_id: str
    message: str

@app.post("/chat")
async def chat_endpoint(msg: UserMessage):
    try:
        # Apply rate limiting
        await rate_limit()
        
        logger.info(f"Processing chat request from user {msg.user_id}: {msg.message[:100]}...")
        
        # Check if message is too long (Gemini has token limits)
        if len(msg.message) > 8000:  # Conservative limit
            return JSONResponse(
                status_code=400, 
                content={"error": "Message too long. Please try a shorter question."}
            )
        
        docs = vectorstore.similarity_search(msg.message, k=1)
        fallback_phrases = [
            "I cannot answer",
            "does not contain",
            "not found",
            "not available",
            "no information",
            "I'm sorry"
        ]
        use_gemini = False
        response = ""
        
        if not docs or docs[0].page_content.strip() == "":
            use_gemini = True
        else:
            # Try to get response from RAG chain with proper async handling
            try:
                # qa_chain.run is synchronous, so we run it in a thread pool
                response = await asyncio.get_event_loop().run_in_executor(
                    None, lambda: qa_chain.run(msg.message)
                )
                if any(phrase.lower() in str(response).lower() for phrase in fallback_phrases):
                    use_gemini = True
            except Exception as e:
                logger.error(f"RAG chain failed: {str(e)}")
                use_gemini = True

        if use_gemini:
            try:
                # Robustly build chat history prompt
                chat_history = memory.load_memory_variables({}).get("chat_history", [])
                history_prompt = ""
                if isinstance(chat_history, list):
                    for turn in chat_history:
                        # LangChain message objects
                        if hasattr(turn, "content"):
                            if getattr(turn, "type", None) == "human":
                                history_prompt += f"User: {turn.content}\n"
                            else:
                                history_prompt += f"AI: {turn.content}\n"
                        # Dict-based messages
                        elif isinstance(turn, dict):
                            if turn.get("type") == "human":
                                history_prompt += f"User: {turn.get('content', '')}\n"
                            else:
                                history_prompt += f"AI: {turn.get('content', '')}\n"
                        # Fallback to string
                        else:
                            history_prompt += str(turn) + "\n"
                elif isinstance(chat_history, str):
                    history_prompt = chat_history
                
                # Limit history length to prevent token overflow
                if len(history_prompt) > 4000:
                    history_prompt = history_prompt[-4000:]
                
                # Add the new user message
                full_prompt = f"{history_prompt}User: {msg.message}\nAI:"
                
                # Create LLM and call it with proper async handling
                gemini_llm = create_gemini_llm()
                response = await asyncio.get_event_loop().run_in_executor(
                    None, lambda: gemini_llm.invoke(full_prompt)
                )
                
                if hasattr(response, 'content'):
                    response = response.content
                elif isinstance(response, dict) and 'content' in response:
                    response = response['content']
                else:
                    response = str(response)
                
                memory.save_context({"input": str(msg.message)}, {"output": str(response)})
                
            except ResourceExhausted:
                return JSONResponse(
                    status_code=429,
                    content={
                        "error": "API quota exceeded. Please wait a moment and try again.",
                        "response": "I'm temporarily overloaded. Please try again in a few seconds."
                    }
                )
            except GoogleAPICallError as e:
                logger.error(f"Google API error: {str(e)}")
                return JSONResponse(
                    status_code=503,
                    content={
                        "error": "Service temporarily unavailable. Please try again.",
                        "response": "I'm having trouble connecting to my knowledge base. Please try again in a moment."
                    }
                )
            except Exception as e:
                logger.error(f"Unexpected error in Gemini fallback: {str(e)}")
                return JSONResponse(
                    status_code=500,
                    content={
                        "error": "Internal server error",
                        "response": "I'm experiencing technical difficulties. Please try again later."
                    }
                )
        else:
            try:
                memory.save_context({"input": str(msg.message)}, {"output": str(response)})
            except Exception as e:
                logger.warning(f"Failed to save context to memory: {str(e)}")
        
        logger.info(f"Successfully processed request for user {msg.user_id}")
        return {"response": str(response)}
        
    except Exception as e:
        logger.error(f"Critical error in chat endpoint: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={
                "error": "Server error",
                "response": "I'm experiencing technical difficulties. Please try again later."
            }
        )

# Set up templates directory
templates = Jinja2Templates(directory="templates")

@app.get("/health")
async def health_check():
    """Health check endpoint to verify server is running"""
    try:
        # Test API key and basic functionality
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            return JSONResponse(
                status_code=500,
                content={"status": "unhealthy", "error": "Missing Google API key"}
            )
        return {"status": "healthy", "message": "Server is running properly"}
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={"status": "unhealthy", "error": str(e)}
        )

# Serve the chat page at root
@app.get("/", response_class=HTMLResponse)
async def serve_chat(request: Request):
    return templates.TemplateResponse("chat.html", {"request": request})

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_image(file_path):
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)

@app.post("/upload")
async def upload_kb_file(file: UploadFile = File(...)):
    if not file.filename:
        return {"error": "No file uploaded."}
    file_ext = file.filename.split(".")[-1].lower()
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    if file_ext == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_ext in ["doc", "docx"]:
        text = extract_text_from_docx(file_path)
    elif file_ext in ["png", "jpg", "jpeg"]:
        text = extract_text_from_image(file_path)
    else:
        return {"error": "Unsupported file type."}
    # Save extracted text to kb.txt (append)
    with open("kb.txt", "a", encoding="utf-8") as kb:
        kb.write(f"\n{text}\n")
    return {"message": "File uploaded and text extracted successfully."}

# To run:
# uvicorn main:app --reload
